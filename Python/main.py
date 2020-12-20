import tkinter
import re
import platform
import pandas as pd
import operator
import os
from bs4 import BeautifulSoup as bs
import numpy as np
import matplotlib.pyplot as plt
from matplotlib import font_manager, rc
from matplotlib.pyplot import figure
figure(num=None, figsize=(27, 10), dpi=80, facecolor='w', edgecolor='k')
plt.rcParams.update({'font.size': 12})
def notice(word):
    window=tkinter.Tk()
    window.title("Lecture : Programming 1 ")
    window.geometry("400x120+100+100")
    window.resizable(False, False)
    label=tkinter.Label(window, text= word )
    label.pack()
    window.mainloop()
notice("! ! ! 특허검색기 입니다 ! ! ! \n이창을 닫으신후, 로딩까지 1-2분정도 소요됩니다.\n 조금만 기다려주세요..!\n검색기 이용 방법 : 숫자, 단어를 입력해주시면 됩니다.\n산업공학과_201611145_박민성\n지도교수님 : 윤장혁 교수님")
print('특허검색기능 로딩중입니다, 로딩까지 1-2분정도 소요됩니다.')

dirname = 'C://특허 검색기 자료 저장'
if not os.path.isdir('C://특허 검색기 자료 저장'):
    os.mkdir('C://특허 검색기 자료 저장')

if platform.system() == 'Windows':
    font_name= font_manager.FontProperties(fname="C:/Windows/Fonts/malgun.ttf").get_name()
    rc('font', family=font_name)
else:
    rc('font', family='AppleGothic')

f = open("개인과제data.txt",'r', encoding='UTF-8')
f = f.read()
f = f.replace("xml version=1.0 encoding=utf-8","xml version=\'1.0\' encoding=\'utf-8\'")
f = f.split('\n')

def makefile(fn):
    save_or_not=input("검색한 내용을 저장하시겠습니까?\n번호를 입력해주세요.\n1) 저장하겠습니다.\n2) 저장하고 싶지 않다.\n")
    if int(save_or_not)==1:
        how = input("출력하고 싶은 파일의 형태를 정해주세요\n1) ##  word 파일로 저장  ##\n 장점 : 깔끔해보인다.\n2) ##  csv 파일로 저장   \
## 로 저장 \n장점 : 다른 프로그램에서 응용하기 쉽고, dict 타입의 자료 저장하기 용이합니다.\n")
        want_name= input("저장하고 싶은 파일명을 적어주세요.\n추천예시)\nex) 발명자_장수길_invetionTitle_자료\n")
        if int(how)==2:
            dataframe = pd.DataFrame(fn)
            a = "C:\\특허 검색기 자료 저장\\"+ want_name+".csv"
            dataframe.to_csv(a, encoding='cp949',header=False, index=False)
            print("C드라이브 -> 특허 검색기 자료 저장 폴더에 저장되었습니다.\n")
        elif int(how)==1:
            from docx import Document
            document = Document()
            document.add_heading(want_name,0)
            p = document.add_paragraph()
            p.add_run("\n")
            for i in fn:
                p.add_run(str(i)+'\n')
            document.save('C:\\특허 검색기 자료 저장\\'+want_name+'.docx')
            print("C->특허 검색기 자료 저장 폴더에 저장되었습니다.\n")
    elif int(save_or_not)==2:
        print("저장하지 않았습니다. 초기화면으로 돌아갑니다.\n")

def list_print_save(list_type):
    number = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n1)리스트 type 으로 화면에서 보기\n2)한줄씩 프린트해서 보기\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
    if number == 1:
        try:
            if len(word_word_save_list) != 0:
                print("*",wantword,"-",wantword2,"의 결과 입니다.\n")
        except:
            pass
        finally:
            print(list_type)
    
    if number == 2:
        try:
            if len(word_word_save_list) != 0:
                print("*",wantword,"-",wantword2,"의 결과 입니다.\n")
        except:
            pass
        finally:
            print("\n")
            for i in range(0,len(list_type)):
                print(list_type[i])
            print("\n")
    makefile(list_type)

def fix(all_list):
    want=input("찾고 싶은 검색어를 입력해주세요\n")
    find_list=[]
    for i in range(0,len(all_list)):
        for j in range(0,len(all_list[i])):
            if want == all_list[i][j]:
                soup = bs(f[i],'xml')
                find_list.append(soup.find('inventionTitle').get_text())
    
    if len(find_list) == 0:
        print("찾으시는 단어 ' "+want+" ' 가 없습니다.")
        new=[]
        for i in all_list:
            for j in i:
                if want in j:
                    new.append(j)
        new=list(set(new))
        
        if len(new)==0:
            hoxy = 2
        
        elif len(new)!=0:
            print('\n')
            print(new)
            hoxy = input("다음은 ' "+want+" ' 를 포함 하는 단어입니다.\n혹시 찾으시는 검색어가 위의 리스트 중에 있나요?\n\n1) 위의 리스트 중에 있다 -> 계속 검색 \n2) 위의 리스트중에 없다 -> 나가기\n")
        
        if int(hoxy) == 1:
            want_agent2= input("위 리스트 중 한 검색어를 다시한번 정확히 입력해주세요\n")
            save_list_2=[]
            
            for i in range(0,len(all_list)):
                for j in range(0,len(all_list[i])):
                    if want_agent2 == all_list[i][j]:
                        soup = bs(f[i],'xml')
                        save_list_2.append(soup.find('inventionTitle').get_text())
            save_list_2=list(set(save_list_2))
            list_print_save(save_list_2)
        
        elif int(hoxy)==2:
            print("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n초기화면으로 돌아갑니다.\n")
    
    elif len(find_list) != 0:
        find_list=list(set(find_list))
        list_print_save(find_list)

def toggae(choose_num,main_list):
    if main_list == agent_toggae_list:
        name = "발명자"
    elif main_list == applicant_toggae_list:
        name = "출원인"
    if choose_num ==1:
        result = dict()
        title = "특허가 많은 상위 10명의" 
        for i in main_list:
            result[i] = main_list.count(i)
        people_num = 10
        final_result= result
    elif choose_num ==2:
        savesave=dict()
        title = "저장한 특허 자료"
        while True:
            
            numnumnum = input("1) 원하는 '"+name+ "' 계속 검색 + 저장\n2) 그만 하고 그래프로 보기\n")
            if int(numnumnum) == 1:
                namename = input("원하시는 '"+name+"'의 이름을 입력해주세요\n")
                result2=dict()
                for i in main_list:
                    result2[i] = main_list.count(i)
                for key,value in result2.items():
                    if key == namename:
                        savesave[key] = value
                print(savesave)
            elif int(numnumnum) == 2:
                break
        
        people_num = len(savesave)
        final_result = savesave
    
    y1_value = []
    x_name =[]
    numnum = sorted(final_result.items(), key = operator.itemgetter(1,0), reverse = True)
    for key,value in numnum:
        y1_value.append(value)
        x_name.append(key)
    a=[]
    b=[]
    for i in range(people_num):
        a.append(y1_value[i])
        b.append(x_name[i])
    x_name = b
    y1_value = a
    n_groups = len(x_name)
    index = np.arange(n_groups)
    plt.bar(index, y1_value, tick_label=x_name, align='center')
    plt.xlabel('특허 개수')
    plt.ylabel( name+' 별 특허 개수')
    plt.title(title+name)
    plt.xlim( -1, n_groups)
    if main_list == agent_toggae_list:
        plt.ylim( 0, 110)
    elif main_list == applicant_toggae_list:
        plt.ylim( 0, 25)
    print(plt.show())
    savenum= int(input("위 그래프를 저장하시겠습니까?\n1) 저장하겠다.\n2) 저장하지 않겠습니다.\n"))
    
    if savenum==1:
        want_name = input("저장하고 싶은 이름을 입력해주세요.\n")
        plt.savefig("C:\\특허 검색기 자료 저장\\"+ want_name+".png",dpi=300)
        print("C드라이브 -> 특허 검색기 자료 저장 폴더에 저장되었습니다.\n")
    elif savenum==2:
        print("초기화면으로 돌아갑니다.\n")

all_agentname=[]
all_applicantname=[]
all_title=[]
agent_toggae_list=[]
applicant_toggae_list=[]
all_content_save=[]
temporary_list=[]

alltitle=[]
view_list=[]
for i in range(len(f)-1):
    soup = bs(f[i],'xml')
    agentInfoArray = soup.findAll('agentInfoArray')
    applicantInfoArray = soup.findAll('applicantInfoArray')
    applicantfider=[]
    agentfinder=[]
    agent_to2list=[]
    applicant_to2list=[]
    all_content = soup.prettify()
    all_content_save.append(all_content)
    titlesplit = soup.find('inventionTitle')
    alltitle.append(titlesplit.get_text())
    for i in range (0,len(agentInfoArray)):
        agentfinder=agentInfoArray[i].findAll('name')
    for i in range(0,len(agentfinder)):
        agent_to2list.append(agentfinder[i].get_text())
        agent_toggae_list.append(agentfinder[i].get_text())
    all_agentname.append(agent_to2list)
    for i in range (0,len(applicantInfoArray)):
        applicantfinder=applicantInfoArray[i].findAll('name')
    for i in range(0,len(applicantfinder)):
        applicant_to2list.append(applicantfinder[i].get_text())
        applicant_toggae_list.append(applicantfinder[i].get_text())
    all_applicantname.append(applicant_to2list)
    titlesplit = soup.find('inventionTitle')
    all_title.append(titlesplit.get_text())

while True:
    try:
        searchnum = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n원하시는 검색기능을 선택해주세요\n1) 특정 출원인->출원한 특허목록출력\n2) 특정 \
키워드포함 특허목록출력\n3) 특정 발명자가 출원한 특허목록 출력\n4) 통계자료 보기\n5) 특정 단어가 들어있는 정보를 모두 출력(아주 자세한 정보를 원하시는분)\n6) 특허 검색기능 종료\n1,2,3,4,5,6 숫자만 입력해주세요.\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
        if searchnum == 1:
            fix(all_applicantname)
        elif searchnum == 2:
            yudonum = 0
            word_save_list = []
            wantword=(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n원하시는 단어를 입력해주세요\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
            for word in all_title:
                if wantword in word:
                    word_save_list.append(word)
            if len(word_save_list) == 0:
                print("* '"+wantword+"' 의 검색 결과가 없습니다. *\n")
                yudonum = 1
            if yudonum == 0:
                researchnum = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n [1번] 현재 검색결과 inventiontitle 출력\n [2번] "" '"+wantword+"' 의 결과 내에서 다시 재검색 하기\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
                if researchnum == 1:
                    word_save_list=list(set(word_save_list))
                    list_print_save(word_save_list)
                elif researchnum ==2:
                    print(word_save_list)
                    word_word_save_list=[]
                    wantword2 = input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n'위의 리스트는 "+wantword+" ' 를 포함하는 특허명입니다.\n재검색을 원하시는 단어를 입력해주세요\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n")
                    for word in word_save_list:
                        if wantword2 in word:
                            word_word_save_list.append(word)
                    if len(word_word_save_list) != 0:
                        
                        word_word_save_list=list(set(word_word_save_list))
                        list_print_save(word_word_save_list)
                    elif len(word_word_save_list) ==0:
                        print("*",wantword,"-",wantword2,"의 결과가 없습니다.\n")
        elif searchnum == 3:
            fix(all_agentname)
                
        elif searchnum==4:
            num = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n1) 발명인에 대한 통계 자료 보기\n2) 출원인에 대한 통계 자료 보기\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
            if num == 1:
                result = dict()
                for i in agent_toggae_list:
                    result[i] = agent_toggae_list.count(i)
                num2 = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n1) 가나다순으로 출력해서 보기\n2) 횟수가 많은 순으로보기\n3) 특정사람 횟수 검색\n4) 특허가 많은 상위 10명 그래프로보기\n5) 인물 검색 후, 횟수 그래프로 출력해서 가져가기\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
                if num2 == 1:
                    ganada = sorted(result.items())
                    print(ganada)
                    makefile(ganada)
                elif num2 ==2:
                    numnum = sorted(result.items(), key = operator.itemgetter(1,0), reverse = True)
                    print(numnum)
                    makefile(numnum)
                elif num2 ==3:
                    want_word2 = input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n원하시는 단어를 입력해주세요\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n")
                    for key,value in result.items():
                        if key == want_word2:
                            print(key,':',value)
                elif num2==4:
                    print(toggae(1,agent_toggae_list))
                elif num2==5:
                    notice("\n원하시는 발명자를 모두 차례대로 입력해주세요\n원하시는 모든 발명자의\n특허의 개수를 그래프로 그려드립니다.")
                    print(toggae(2,agent_toggae_list))
            elif num==2:
                result = dict()
                for i in applicant_toggae_list:
                    result[i] = applicant_toggae_list.count(i)
                num2 = int(input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n1) 가나다순으로 출력해서 보기\n2) 횟수가 많은 순으로보기\n3) 특정 출원인 횟수 검색\n4) 특허가 많은 상위 10명 그래프로보기\n5) 특정 출원인 검색 후, 횟수 그래프로 출력해서 가져가기\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n"))
                if num2 == 1:
                    ganada = sorted(result.items())
                    print(ganada)
                    makefile(ganada)
                elif num2 ==2:
                    numnum = sorted(result.items(), key = operator.itemgetter(1,0), reverse = True)
                    print(numnum)
                    makefile(numnum)
                elif num2 ==3:
                    want_word2 = input("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n원하시는 단어를 입력해주세요\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n")
                    for key,value in result.items():
                        if key == want_word2:
                            print(key,':',value)
                elif num2==4:
                    print(toggae(1,applicant_toggae_list))
                elif num2==5:
                    notice("\n원하시는 출원인을 모두 차례대로 입력해주세요\n원하시는 모든 출원인의\n특허의 개수를 그래프로 그려드립니다.")
                    print(toggae(2,applicant_toggae_list))

        elif searchnum==5:
            word = input("찾고 싶은 단어가 무엇인가요??\n")
            for i in range(0,len(all_content_save)):
                if word in all_content_save[i]:
                    a=all_content_save[i]
                    text = re.sub('</.+?>', '',a, 0, re.I|re.S)
                    text2 = re.sub('<.+?:', '',text, 0, re.I|re.S)
                    text3 = re.sub('type="kr.or.kipris.plus.webservice.services.patentbean.+?>', '',text2, 0, re.I|re.S)
                    textxx = text3
                    text3 = re.sub('>','',text3)
                    temporary_list.append(text3)
                    view_list.append(alltitle[i])
            view_list=list(set(view_list))
            print(view_list)
            
            accurate_num=input("' "+word+" '가 들어 있는 inventiontitle의 모든 목록 입니다.\n 어떤 발명에 대한 정보를 출력할까요?\n위의 목록에서 보고싶은 발명 하나를 정확히 입력해주세요\n")
            text3=temporary_list[view_list.index(accurate_num)]
            text3=text3.split('\n')
            mid=[]
            last=[]
            for i in range(len(text3)):
                if text3[i].strip() != '':
                    mid.append(text3[i].strip())
            for i in range(4,len(mid)-1):
                    print(mid[i])
                    last.append(mid[i])
            
            print(makefile(last))
        elif searchnum==6:
            notice("검색기를 이용해주셔서 대단히 감사합니다.\n건국대학교\n산업공학과_201611145_박민성\nLECTURE : Programming 1\n지도 교수님 : 윤장혁 교수님\n도움을 주신분 : 오승현 조교님, 이지호 조교님 ")
            break
    except  ValueError:
        print("ㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ\n 숫자만 입력해주세요.\n 정확한 숫자를 입력해주세요\n 초기화면으로 돌아갑니다.\nㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡㅡ")